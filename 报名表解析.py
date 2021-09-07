import time

import openpyxl
from docx import Document
import os
import shutil
from icecream import ic
import zipfile


def SignUpDocx(TITLE, i, path, primer_PATH):  # 文件生成
    document = Document()
    document.add_heading("报名表", level=0)
    for count in range(11):
        document.add_heading(TITLE[count], level=1, )
        document.add_paragraph(i[count], style="Intense Quote")
    document.save(primer_PATH + path + i[0] + ".docx")  # 可以覆盖保存，PATH是根目录


def CreateFolder(path, name_list):  # 文件夹生成
    for i in name_list:  # ['第一志愿/','第二志愿/']:
        isExists = os.path.exists(path + str(i))
        if not isExists:
            os.makedirs(path + str(i))
            print("{} 目录创建成功".format(i))
        else:
            print("{} 目录已经存在，文件删除出错，请检查".format(i))
            break


def RemoveFolder(path, folder_name):
    if os.path.exists(path + folder_name):
        shutil.rmtree(path + folder_name)


def DuplicateRemoval(prime_list):  # 根据学号和手机去重
    for i in prime_list:
        for j in prime_list[:prime_list.index(i)]:
            if j[2] == i[2] or j[3] == i[3] or j[4] == i[4]:
                print("+{:6}已被去重".format(i[1]))
                prime_list.remove(j)
    return prime_list


#
#
def Shunt(wish, TITLE, sign_up_list, primer_PATH):  # 分流
    for i in sign_up_list[1:]:
        SignUpDocx(TITLE, i, wish + i[TITLE.index(wish[:-1])] + "/", primer_PATH)  # wish[:-1]是去掉/


def load_excel(excel_file):
    wb = openpyxl.load_workbook(excel_file)
    ws = wb[wb.sheetnames[0]]
    # for i in ws["C":"M"]:
    count = 0
    rows_list = []
    for row in ws.values:
        row = row[2:13]
        if count == 0:
            # ic(row)
            row = [i.split("_")[1] for i in row]
            # ic(row)
            count += 1
        rows_list.append(row)
    os.remove(excel_file)
    for i in os.listdir("."):
        if ".xlsx" in i:
            os.remove(i)
    return rows_list


def get_all_file_paths(directory):
    # 初始化文件路径列表
    file_paths = []
    for root, directories, files in os.walk(directory):
        for filename in files:
            # 连接字符串形成完整的路径
            filepath = os.path.join(root, filename)
            file_paths.append(filepath)

    # 返回所有文件路径
    return file_paths


def file2zip(name):
    for i in os.listdir("."):
        if ".zip" in i:
            os.remove(i)
    now_time = time.strftime(" %m-%d(%H：%M)", time.localtime())
    # zip_file = zipfile.ZipFile(name + now_time + '.zip', 'w')
    # # 把zfile整个目录下所有内容，压缩为new.zip文件
    # zip_file.write('zfile', compress_type=zipfile.ZIP_DEFLATED)
    # # 把c.txt文件压缩成一个压缩文件
    # # zip_file.write('c.txt',compress_type=zipfile.ZIP_DEFLATED)
    # zip_file.close()
    with zipfile.ZipFile(name + now_time + '.zip', 'w') as my_zip:
        # my_zip.write('./报名表/2021年学生社团联合会志愿者电子报名表.xlsx')
        # my_zip.write("./报名表")
        for file in get_all_file_paths("./报名表"):
            my_zip.write(file)

def write_excel(signUp_list, excel_name):
    bk1 = openpyxl.Workbook()
    sheet1 = bk1.active
    for i in signUp_list:
        sheet1.append(i)
    bk1.save(excel_name)


if __name__ == "__main__":
    PATH = "./报名表/"
    # f = open(PATH + "2021年学生社团联合会志愿者电子报名表.csv", "r", encoding="UTF-8")
    for i in os.listdir("."):
        if ".xlsx" in i:
            excel_reader = load_excel(i)
            sign_up_list = tuple(DuplicateRemoval(list(excel_reader)))  # 去重
            write_excel(sign_up_list, "./报名表/2021年学生社团联合会志愿者电子报名表.xlsx")
            TITLE = tuple(sign_up_list[0])

            RemoveFolder(PATH, '第一志愿/')
            RemoveFolder(PATH, '第二志愿/')  # 删除非空文件夹，上同
            CreateFolder(PATH, ['第一志愿/', '第二志愿/'])  # 删除之前的数据，已去掉重复填报者

            for wish in ['第一志愿/', '第二志愿/']:
                CreateFolder(PATH + wish, ["宣传媒体中心", "办公室", "社团活动管理部", "就业部", "财务与监察部", "社团事务部", "阳光服务部", "自律会"])
                Shunt(wish, TITLE, sign_up_list, PATH)
            file2zip("报名表")
            break
